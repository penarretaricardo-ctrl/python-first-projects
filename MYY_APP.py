
from docx import Document
from docx.shared import Inches
import pyttsx3

# --- TTS: init once ---
engine = pyttsx3.init()            # reuse the same engine
engine.setProperty('rate', 180)    # optional: adjust speed
engine.setProperty('volume', 1.0)  # optional

def speak(text):
    engine.say(text)
    engine.runAndWait()

document = Document()

# picture (won't crash silently if missing)
try:
    document.add_picture('profile_picture.jpg', width=Inches(2.0))
except Exception as e:
    print("Skipping picture:", e)

# ---- Inputs + speech ----
name = input('What is your name? ')
speak(f"Hello {name}, welcome to the CV generator.")
speak("Let's get started with your CV.")

phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

document.add_paragraph(f"{name} | {phone_number} | {email}")

document.add_heading('About Me')
document.add_paragraph(input('Tell me about yourself: '))

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
start_date = input('Enter start date: ')
end_date = input('Enter end date: ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + ' - ' + end_date + '\n').italic = True
experience_details = input(f'Describe your experience at {company}: ')
p.add_run(experience_details)

# More experiences
while True:
    more = input('Do you have any extra experiences? Yes or No? ').strip().lower()
    if more == 'yes':
        p = document.add_paragraph()
        company = input('Enter company: ')
        start_date = input('Enter start date: ')
        end_date = input('Enter end date: ')
        p.add_run(company + ' ').bold = True
        p.add_run(start_date + ' - ' + end_date + '\n').italic = True
        experience_details = input(f'Describe your experience at {company}: ')
        p.add_run(experience_details)
    else:
        break

document.add_heading('Skills')
skill_paragraph = document.add_paragraph(input('Enter skill: '))
skill_paragraph.style = 'List Bullet'
while True:
    more = input('Do you have any extra skills? Yes or No? ').strip().lower()
    if more == 'yes':
        skill_paragraph = document.add_paragraph(input('Enter one of your skills: '))
        skill_paragraph.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
footer.paragraphs[0].text = "This CV was generated using Python"

document.save('personalcv.docx')
speak("Your CV has been generated.")
